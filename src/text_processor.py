import io
import os
import uuid
from docx import Document
import docx2txt
import fitz
import pypandoc
import streamlit as st
import openai
from google import genai
from config import OPENAI_API_KEY, GEMINI_API_KEY, XAI_API_KEY
from src.prompts import GENERAL_SUMMARY_PROMPT

# Model configuration
MODEL_CONFIG = {
    "gpt-4.1": {
        "client_type": "openai",
        "base_url": "https://api.openai.com/v1",
        "api_key": OPENAI_API_KEY
    },
    "o4-mini": {
        "client_type": "openai",
        "base_url": "https://api.openai.com/v1",
        "api_key": OPENAI_API_KEY
    },
    "grok-3": {
        "client_type": "openai",
        "base_url": "https://api.x.ai/v1",
        "api_key": XAI_API_KEY
    },
    "grok-3-mini": {
        "client_type": "openai",
        "base_url": "https://api.x.ai/v1",
        "api_key": XAI_API_KEY
    },
    "gemini-2.0-flash": {
        "client_type": "gemini",
        "api_key": GEMINI_API_KEY
    },
    "gemini-2.5-pro-exp-03-25": {
        "client_type": "gemini",
        "api_key": GEMINI_API_KEY
    },
    "gemini-1.5-pro-latest": {
        "client_type": "gemini",
        "api_key": GEMINI_API_KEY
    }
}

def summarize_transcription(transcription_input, model="gemini-2.0-flash", custom_prompt=None, enable_reasoning=False):
    """
    Generate a summary from transcription data using the specified model.
    
    Args:
        transcription_input: The transcription data as JSON or plain text
        model: The model to use for summarization (e.g., gpt-4.1, grok-3, gemini-2.5-pro)
        custom_prompt: Optional custom system prompt for summarization
        enable_reasoning: Whether to include step-by-step reasoning (for supported models)
        
    Returns:
        Tuple of (summary text, reasoning text)
    """
    config = MODEL_CONFIG[model]
    full_prompt = custom_prompt if custom_prompt else GENERAL_SUMMARY_PROMPT
    
    # Add reasoning instruction for supported models
    if enable_reasoning and model in ["grok-3-mini", "gemini-2.5-pro-exp-03-25", "o4-mini"]:
        full_prompt += "\n\nProvide a step-by-step reasoning process before generating the final summary."

    # Handle different input types
    if isinstance(transcription_input, list):
        # JSON format from direct transcription
        transcription_text = "\n".join(
            f"[{entry['timestamp']}] {entry['speaker']}: {entry['text']}"
            for entry in transcription_input
        )
    elif isinstance(transcription_input, str):
        # Already text from DOCX/PDF extraction
        transcription_text = transcription_input
    else:
        raise ValueError("Unsupported transcription input format")

    try:
        if config["client_type"] == "openai" and model == "o4-mini":
            # Initialize OpenAI client
            client = openai.OpenAI(
                api_key=config["api_key"],
                base_url=config["base_url"]
            )
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": full_prompt},
                    {"role": "user", "content": transcription_text}
                ],
                max_completion_tokens=20000
                )
            content = response.choices[0].message.content
        elif config["client_type"] == "openai":
            # Initialize OpenAI client for o4-mini
            client = openai.OpenAI(
                api_key=config["api_key"],
                base_url=config["base_url"]
            )
            response = client.chat.completions.create(
                model=model,
                messages=[
                    {"role": "system", "content": full_prompt},
                    {"role": "user", "content": transcription_text}
                ],
                max_tokens=20000
            )
            content = response.choices[0].message.content
        else:  # Gemini
            # Initialize Gemini client
            client = genai.Client(api_key=config["api_key"])
            response = client.models.generate_content(
                model=model,
                contents=[full_prompt, transcription_text]
            )
            content = response.text

        # Extract reasoning if enabled
        reasoning = ""
        if enable_reasoning and model in ["grok-3-mini", "gemini-2.5-pro-exp-03-25", "o4-mini"]:
            # Assume reasoning is before a separator or the summary
            parts = content.split("\n\n---\n\n", 1)  # Adjust based on API response format
            reasoning = parts[0].strip() if len(parts) > 1 else ""
            content = parts[-1].strip() if len(parts) > 1 else content

        return content, reasoning

    except Exception as e:
        raise Exception(f"Error generating summary with {model}: {str(e)}")

def export_transcription_to_docx(transcription_json, output_folder="transcripts", file_name="transcript.docx"):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    output_path = os.path.join(output_folder, file_name)
    if os.path.exists(output_path):
        base, ext = os.path.splitext(file_name)
        file_name = f"{base}_{uuid.uuid4().hex[:8]}{ext}"
        output_path = os.path.join(output_folder, file_name)
    doc = Document()
    doc.add_heading("Meeting Transcription", level=1)
    for entry in transcription_json:
        doc.add_paragraph(f"[{entry['timestamp']}] {entry['speaker']}: {entry['text']}")
    doc.save(output_path)
    return output_path

def export_summary_to_docx(summary, output_folder="transcripts", file_name="summary.docx"):
    if not os.path.exists(output_folder):
        os.makedirs(output_folder)
    output_path = os.path.join(output_folder, file_name)
    if os.path.exists(output_path):
        base, ext = os.path.splitext(file_name)
        file_name = f"{base}_{uuid.uuid4().hex[:8]}{ext}"
        output_path = os.path.join(output_folder, file_name)

    # Prepare the Markdown content
    markdown_content = "# Meeting Summary\n\n" + summary

    try:
        # Use pypandoc to convert Markdown to DOCX
        pypandoc.convert_text(
            markdown_content,
            'docx',
            format='md',
            outputfile=output_path
        )
        print(f"Converted Markdown to DOCX: {output_path}")
    except OSError as e:
        raise Exception(f"Pandoc conversion failed: {str(e)}. Please ensure Pandoc is installed on your system.")

    return output_path

def extract_text_from_docx(file_bytes):
    """Extract text from a DOCX file."""
    try:
        text = docx2txt.process(io.BytesIO(file_bytes))
        return text
    except Exception as e:
        st.error(f"Error extracting text from DOCX: {e}")
        return None

def extract_text_from_pdf(file_bytes):
    """Extract text from a PDF file."""
    try:
        with fitz.open(stream=file_bytes, filetype="pdf") as doc:
            text = ""
            for page in doc:
                text += page.get_text("text") + "\n"
            return text
    except Exception as e:
        st.error(f"Error extracting text from PDF: {e}")
        return None